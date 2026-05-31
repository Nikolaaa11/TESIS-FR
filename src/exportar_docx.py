"""
exportar_docx.py
----------------
Genera docs/Tesis_USS.docx con formato de TESIS DE MAGÍSTER:
  - Portada académica (universidad, facultad, programa, título, autor, guía, ciudad/año).
  - Índice automático (campo TOC, se actualiza en Word con F9).
  - Encabezados numerados, cuerpo justificado Times New Roman 12, interlineado 1.5.
  - Tablas con estilo, figuras con numeración de caption.
  - Encabezado de página (título corto) y pie con número de página.

Parser markdown ligero sobre docs/tesis.md. Si docs/ está bloqueado (Word abierto)
guarda igualmente la copia en web/assets/docs/.
"""
import os, sys, re, glob
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.path.append(os.path.dirname(__file__))
import config as C

NAVY = RGBColor(0x1B, 0x2A, 0x41)
COPPER = RGBColor(0xC2, 0x70, 0x3D)
GRAY = RGBColor(0x55, 0x55, 0x55)
BODY_FONT = "Times New Roman"
HEAD_FONT = "Calibri"


# ----------------------- helpers de bajo nivel -----------------------
def _set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor); tcPr.append(shd)


def _field(paragraph, instr):
    """Inserta un campo de Word (TOC, PAGE, etc.)."""
    r = paragraph.add_run()
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    instr_el = OxmlElement("w:instrText"); instr_el.set(qn("xml:space"), "preserve"); instr_el.text = instr
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = "Actualice este campo en Word (F9)"
    f3 = OxmlElement("w:fldChar"); f3.set(qn("w:fldCharType"), "end")
    r._r.append(f1); r._r.append(instr_el); r._r.append(f2); r._r.append(t); r._r.append(f3)


def _page_number(paragraph):
    run = paragraph.add_run()
    for typ, txt in (("begin", None), (None, "PAGE"), ("separate", None), (None, "1"), ("end", None)):
        if typ in ("begin", "separate", "end"):
            fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), typ); run._r.append(fc)
        elif txt == "PAGE":
            it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = " PAGE "; run._r.append(it)


# ----------------------- estilos del documento -----------------------
def _estilos(doc):
    st = doc.styles["Normal"]
    st.font.name = BODY_FONT; st.font.size = Pt(12)
    pf = st.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_after = Pt(4)
    pf.first_line_indent = Cm(0.8)   # sangría de primera línea (estándar de tesis)
    for i, sz in [(1, 16), (2, 14), (3, 12)]:
        h = doc.styles[f"Heading {i}"]
        h.font.name = HEAD_FONT; h.font.size = Pt(sz); h.font.color.rgb = NAVY; h.font.bold = True
        h.paragraph_format.space_before = Pt(14 if i == 1 else 10); h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
    # márgenes
    for s in doc.sections:
        s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5)
        s.left_margin = Cm(3); s.right_margin = Cm(2.5)


def _header_footer(doc):
    sec = doc.sections[-1]
    sec.different_first_page_header_footer = True  # portada sin encabezado/número
    hdr = sec.header.paragraphs[0]
    hdr.text = "Impacto macro-financiero en la valoración del cobre en Chile"
    hdr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in hdr.runs:
        r.font.size = Pt(8); r.font.color.rgb = GRAY; r.italic = True
    ftr = sec.footer.paragraphs[0]; ftr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _page_number(ftr)
    for r in ftr.runs:
        r.font.size = Pt(9); r.font.color.rgb = GRAY


# ----------------------- portada e índice -----------------------
def portada(doc):
    def cl(txt, sz, bold=False, color=None, after=6, italic=False):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(after); p.paragraph_format.line_spacing = 1.15
        r = p.add_run(txt); r.font.name = HEAD_FONT; r.font.size = Pt(sz); r.bold = bold; r.italic = italic
        if color: r.font.color.rgb = color
        return p
    doc.add_paragraph()
    # logo USS si existe (assets/logo_uss.png)
    logo = C.ROOT / "assets" / "logo_uss.png"
    if logo.exists():
        try:
            pL = doc.add_paragraph(); pL.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pL.add_run().add_picture(str(logo), width=Cm(3.2))
        except Exception:
            pass
    cl("UNIVERSIDAD SAN SEBASTIÁN", 18, True, NAVY, 2)
    cl("Facultad de Economía y Negocios", 12, False, GRAY, 2)
    cl("Magíster en Data Science", 13, True, None, 24)
    for _ in range(2): doc.add_paragraph()
    cl("Impacto de las variables macroeconómicas globales y financieras en la "
       "valoración bursátil del sector de minería de cobre en Chile", 18, True, NAVY, 6)
    cl("Un análisis econométrico de series de tiempo y panel, 2004–2026", 12, False, GRAY, 30, italic=True)
    for _ in range(4): doc.add_paragraph()
    cl("Tesis para optar al grado de Magíster en Data Science", 12, False, None, 24, italic=True)
    cl("Autor: ________________________", 12, False, None, 4)
    cl("Profesor guía: ________________________", 12, False, None, 24)
    cl("Santiago de Chile · 2026", 12, True, NAVY, 4)
    doc.add_page_break()


def indice(doc):
    h = doc.add_heading("Índice", level=1)
    p = doc.add_paragraph()
    _field(p, 'TOC \\o "1-3" \\h \\z \\u')
    doc.add_page_break()


# ----------------------- parser markdown -----------------------
def _clean(t):
    t = re.sub(r"\*\*(.+?)\*\*", r"\1", t); t = re.sub(r"`(.+?)`", r"\1", t)
    t = re.sub(r"\$\$(.+?)\$\$", r"\1", t); t = re.sub(r"\\\((.+?)\\\)", r"\1", t)
    t = re.sub(r"\*(.+?)\*", r"\1", t)
    t = re.sub(r"\[(.+?)\]\((.+?)\)", r"\1 (\2)", t)
    return t


def _runs_bold(par, texto):
    for p in re.split(r"(\*\*.+?\*\*)", texto):
        if p.startswith("**") and p.endswith("**"):
            par.add_run(_clean(p)).bold = True
        else:
            par.add_run(_clean(p))


def _tabla(doc, lineas, ntab=0, titulo=""):
    filas = []
    for l in lineas:
        if re.match(r"^\s*\|[\s:\-\|]+\|\s*$", l): continue
        filas.append([c.strip() for c in l.strip().strip("|").split("|")])
    if not filas: return
    if ntab:
        tit = re.sub(r"^\d+(\.\d+)*\.?\s+", "", titulo).strip()
        cap = doc.add_paragraph(); cap.paragraph_format.first_line_indent = Cm(0)
        cap.paragraph_format.space_before = Pt(6); cap.paragraph_format.space_after = Pt(3)
        r = cap.add_run(f"Tabla {ntab}. {_clean(tit)}")
        r.bold = True; r.font.size = Pt(9.5); r.font.name = HEAD_FONT; r.font.color.rgb = NAVY
    ncol = max(len(f) for f in filas)
    t = doc.add_table(rows=0, cols=ncol); t.style = "Light Grid Accent 1"
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, fila in enumerate(filas):
        fila += [""] * (ncol - len(fila)); cells = t.add_row().cells
        for j, val in enumerate(fila):
            cells[j].text = ""
            par = cells[j].paragraphs[0]; par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = par.add_run(_clean(val)); run.font.size = Pt(9.5); run.font.name = BODY_FONT
            if i == 0:
                run.bold = True; run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                _set_cell_bg(cells[j], "1B2A41")
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _fig(doc, archivo, nfig, leyenda):
    p = C.FIG / archivo
    if not p.exists(): return
    try:
        pimg = doc.add_paragraph(); pimg.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pimg.paragraph_format.first_line_indent = Cm(0)
        pimg.add_run().add_picture(str(p), width=Cm(14.0))
        cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.first_line_indent = Cm(0); cap.paragraph_format.space_after = Pt(10)
        r = cap.add_run(f"Figura {nfig}. {_clean(leyenda)}"); r.italic = True; r.font.size = Pt(9); r.font.color.rgb = GRAY
    except Exception:
        pass


def _tabla_csv(doc, nombre, ntab, leyenda, maxrows=14):
    p = C.TAB / nombre
    if not p.exists(): return
    df = pd.read_csv(p).head(maxrows)
    filas = [list(df.columns)] + df.astype(str).values.tolist()
    lineas = ["| " + " | ".join(str(c) for c in f) + " |" for f in filas]
    _tabla(doc, lineas, ntab, leyenda)


def convertir(doc, md):
    lineas = md.split("\n"); i = 0; titulo_omitido = False; ntab = 0; nfig = 0; ultimo_tit = "Resultados"
    while i < len(lineas):
        s = lineas[i].strip()
        if not s: i += 1; continue
        mf = re.match(r"^\[\[FIG:\s*([^|\]]+?)\s*\|\s*(.+?)\s*\]\]$", s)
        if mf:
            nfig += 1; _fig(doc, mf.group(1).strip(), nfig, mf.group(2).strip()); i += 1; continue
        mc = re.match(r"^\[\[CSV:\s*([^|\]]+?)\s*\|\s*(.+?)\s*\]\]$", s)
        if mc:
            ntab += 1; _tabla_csv(doc, mc.group(1).strip(), ntab, mc.group(2).strip()); i += 1; continue
        if s.startswith("# ") and not titulo_omitido:
            titulo_omitido = True; i += 1; continue  # título ya está en la portada
        if s.startswith("|"):
            blk = []
            while i < len(lineas) and lineas[i].strip().startswith("|"):
                blk.append(lineas[i]); i += 1
            ntab += 1; _tabla(doc, blk, ntab, ultimo_tit); continue
        if s.startswith("> "):
            p = doc.add_paragraph(style="Intense Quote"); _runs_bold(p, s[2:]); i += 1; continue
        if s.startswith("### "): ultimo_tit = _clean(s[4:]); doc.add_heading(ultimo_tit, level=3)
        elif s.startswith("## "): ultimo_tit = _clean(s[3:]); doc.add_heading(ultimo_tit, level=2)
        elif s.startswith("# "):
            doc.add_heading(_clean(s[2:]), level=1)
        elif re.match(r"^[-*] ", s):
            p = doc.add_paragraph(style="List Bullet"); _runs_bold(p, s[2:])
        elif re.match(r"^\d+\.\s", s):
            p = doc.add_paragraph(style="List Number"); _runs_bold(p, re.sub(r"^\d+\.\s", "", s))
        elif set(s) <= set("-") and len(s) >= 3: pass
        else:
            p = doc.add_paragraph(); _runs_bold(p, s)
        i += 1


def anexo_figuras(doc):
    doc.add_page_break(); doc.add_heading("Anexo A — Figuras del análisis", level=1)
    figs = sorted(glob.glob(str(C.FIG / "*.png")))
    nombres = {
        "precios_normalizados": "Evolución de precios normalizados (base 100) y precio del cobre",
        "irf_ANTO_L": "Función impulso-respuesta: Antofagasta ante shock del cobre",
        "irf_PUCOBRE_SN": "Función impulso-respuesta: Pucobre (respuesta diferida)",
        "heatmap_correlaciones": "Matriz de correlaciones de retornos",
        "iliquidez_vs_beta": "Iliquidez vs transmisión contemporánea del cobre",
    }
    n = 1
    for f in figs:
        key = os.path.basename(f).replace(".png", "")
        cap = nombres.get(key, key.replace("_", " "))
        try:
            doc.add_picture(f, width=Inches(5.8))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            pc = doc.add_paragraph(); pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = pc.add_run(f"Figura A.{n}. {cap}"); r.italic = True; r.font.size = Pt(9); r.font.color.rgb = GRAY
            n += 1
        except Exception as e:
            doc.add_paragraph(f"[no se pudo insertar {key}: {e}]")


def construir():
    md = (C.ROOT / "docs" / "tesis.md").read_text(encoding="utf-8")
    doc = Document(); _estilos(doc)
    portada(doc); indice(doc); convertir(doc, md); anexo_figuras(doc)
    _header_footer(doc)
    # guardar (web siempre; docs si no está bloqueado)
    web = C.ROOT / "web" / "assets" / "docs" / "Tesis_USS.docx"
    web.parent.mkdir(parents=True, exist_ok=True); doc.save(web)
    print(f"OK {web}")
    try:
        doc.save(C.ROOT / "docs" / "Tesis_USS.docx"); print("OK docs/Tesis_USS.docx")
    except PermissionError:
        print("AVISO docs/Tesis_USS.docx bloqueado (Word abierto). Solo se actualizó la copia web; ciérralo y re-corre el script.")


if __name__ == "__main__":
    construir()
