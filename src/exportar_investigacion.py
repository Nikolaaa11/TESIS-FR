"""
exportar_investigacion.py
-------------------------
Convierte los documentos de investigacion/ en un DOSSIER único, en PDF y Word, con
estilo "Apple": fuente Segoe UI (cercana a San Francisco), fondo blanco, títulos
grandes y limpios, color de acento, mucho aire. Une líneas en párrafos correctos.

Salida: docs/Dossier_Investigacion.pdf y docs/Dossier_Investigacion.docx
        (+ copias en web/assets/docs/)
"""
import os, sys, re, glob
import pandas as pd
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from reportlab.lib.styles import ParagraphStyle
from reportlab.platypus import (BaseDocTemplate, PageTemplate, Frame, Paragraph, Spacer,
                                Table, TableStyle, PageBreak, KeepTogether)
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

sys.path.append(os.path.dirname(__file__))
import config as C

INK = colors.HexColor("#1D1D1F"); SUB = colors.HexColor("#6E6E73")
NAVY = colors.HexColor("#1B2A41"); COPPER = colors.HexColor("#C2703D")
LINE = colors.HexColor("#E3E3E6"); SOFT = colors.HexColor("#F5F5F7")
WF = r"C:\Windows\Fonts"

DOCS = [
    ("investigacion/02_historia_del_cobre.md", "Historia del cobre"),
    ("investigacion/01_estado_del_arte.md", "Estado del arte"),
    ("investigacion/03_bibliografia_anotada.md", "Bibliografía anotada"),
    ("investigacion/04_bitacora_investigacion.md", "Bitácora de la investigación"),
]


def _reg():
    fam = {"r": "Helvetica", "b": "Helvetica-Bold", "i": "Helvetica-Oblique", "sb": "Helvetica-Bold"}
    try:
        for n, f in {("SUI", "segoeui.ttf"), ("SUI-B", "segoeuib.ttf"),
                     ("SUI-SB", "seguisb.ttf"), ("SUI-I", "segoeuii.ttf")}:
            pdfmetrics.registerFont(TTFont(n, os.path.join(WF, f)))
        from reportlab.pdfbase.pdfmetrics import registerFontFamily
        registerFontFamily("SUI", normal="SUI", bold="SUI-B", italic="SUI-I", boldItalic="SUI-B")
        fam = {"r": "SUI", "b": "SUI-B", "i": "SUI-I", "sb": "SUI-SB"}
    except Exception:
        pass
    return fam
F = _reg()


def _clean(t):
    t = "" if t is None else str(t)
    t = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", t)
    t = re.sub(r"`(.+?)`", r"\1", t)
    t = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", r"<i>\1</i>", t)
    t = re.sub(r"\[(.+?)\]\((.+?)\)", r'<font color="#0066CC">\1</font>', t)
    t = t.replace("&", "&amp;")
    for tag in ("b", "i", "font", "/font"):
        t = t.replace(f"&lt;{tag}&gt;", f"<{tag}>")
    t = t.replace('&lt;font color="#0066CC"&gt;', '<font color="#0066CC">').replace("&lt;/font&gt;", "</font>")
    return t


def estilos():
    body = ParagraphStyle("b", fontName=F["r"], fontSize=10.5, leading=16, alignment=TA_LEFT,
                          spaceAfter=9, textColor=INK)
    h1 = ParagraphStyle("h1", fontName=F["sb"], fontSize=23, leading=27, textColor=NAVY, spaceBefore=6, spaceAfter=14)
    h2 = ParagraphStyle("h2", fontName=F["sb"], fontSize=15, leading=20, textColor=INK, spaceBefore=16, spaceAfter=7)
    h3 = ParagraphStyle("h3", fontName=F["sb"], fontSize=12, leading=16, textColor=COPPER, spaceBefore=12, spaceAfter=5)
    bullet = ParagraphStyle("bu", parent=body, leftIndent=16, bulletIndent=4, spaceAfter=6)
    quote = ParagraphStyle("q", parent=body, leftIndent=14, fontName=F["i"], textColor=SUB, fontSize=10, leading=15)
    eyebrow = ParagraphStyle("ey", fontName=F["sb"], fontSize=10, leading=13, textColor=COPPER, spaceAfter=2)
    return dict(body=body, h1=h1, h2=h2, h3=h3, bullet=bullet, quote=quote, eyebrow=eyebrow)


def tabla(lineas, S):
    filas = [[c.strip() for c in l.strip().strip("|").split("|")]
             for l in lineas if not re.match(r"^\s*\|[\s:\-\|]+\|\s*$", l)]
    if not filas: return []
    ncol = max(len(f) for f in filas)
    cH = ParagraphStyle("cH", fontName=F["sb"], fontSize=9, leading=12, textColor=colors.white)
    cc = ParagraphStyle("cc", fontName=F["r"], fontSize=9, leading=12, textColor=INK)
    data = [[Paragraph(_clean(c), cH if r == 0 else cc) for c in (f + [""] * (ncol - len(f)))]
            for r, f in enumerate(filas)]
    avail = A4[0] - 4.4*cm
    t = Table(data, colWidths=[avail/ncol]*ncol, hAlign="LEFT")
    t.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, 0), NAVY),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, SOFT]),
        ("GRID", (0, 0), (-1, -1), 0.4, LINE), ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 5), ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("LEFTPADDING", (0, 0), (-1, -1), 7)]))
    return [Spacer(1, 4), t, Spacer(1, 10)]


def parse_md(md, S, skip_h1=True):
    flow = []; lineas = md.split("\n"); i = 0; buf = []; first_h1_skipped = not skip_h1
    def flush():
        if buf:
            flow.append(Paragraph(_clean(" ".join(x.strip() for x in buf)), S["body"])); buf.clear()
    while i < len(lineas):
        s = lineas[i].strip()
        if not s: flush(); i += 1; continue
        if s.startswith("|"):
            flush(); blk = []
            while i < len(lineas) and lineas[i].strip().startswith("|"):
                blk.append(lineas[i]); i += 1
            flow += tabla(blk, S); continue
        if s.startswith("# "):
            flush()
            if not first_h1_skipped:
                first_h1_skipped = True; i += 1; continue
            flow.append(Paragraph(_clean(s[2:]), S["h2"]))
        elif s.startswith("### "): flush(); flow.append(Paragraph(_clean(s[4:]), S["h3"]))
        elif s.startswith("## "): flush(); flow.append(Paragraph(_clean(s[3:]), S["h2"]))
        elif re.match(r"^[-*] ", s): flush(); flow.append(Paragraph("•&nbsp;&nbsp;" + _clean(s[2:]), S["bullet"]))
        elif s.startswith("> "): flush(); flow.append(Paragraph(_clean(s[2:]), S["quote"]))
        elif set(s) <= set("-") and len(s) >= 3: flush()
        else: buf.append(s)
        i += 1
    flush()
    return flow


def _cover(canvas, doc):
    canvas.saveState(); W, H = A4
    canvas.setFillColor(NAVY); canvas.rect(0, H-7.0*cm, W, 7.0*cm, fill=1, stroke=0)
    canvas.setFillColor(colors.white); canvas.setFont(F["sb"], 30)
    canvas.drawString(2.2*cm, H-3.4*cm, "Dossier de Investigación")
    canvas.setFont(F["r"], 13); canvas.setFillColor(colors.HexColor("#CFD6E2"))
    canvas.drawString(2.2*cm, H-4.3*cm, "Estado del arte · Historia del cobre · Bibliografía · Bitácora")
    canvas.setFillColor(COPPER); canvas.rect(2.2*cm, H-4.85*cm, 3.2*cm, 0.12*cm, fill=1, stroke=0)
    canvas.setFillColor(INK); canvas.setFont(F["sb"], 14)
    canvas.drawString(2.2*cm, H-9.5*cm, "Impacto de las variables macroeconómicas y financieras")
    canvas.drawString(2.2*cm, H-10.2*cm, "en la valoración bursátil del cobre en Chile")
    canvas.setFont(F["r"], 11.5); canvas.setFillColor(SUB)
    canvas.drawString(2.2*cm, H-11.2*cm, "Material de investigación · Magíster en Data Science · Universidad San Sebastián")
    canvas.setFont(F["r"], 10); canvas.setFillColor(SUB)
    canvas.drawString(2.2*cm, 2.0*cm, "Documento de apoyo a la tesis · Fuentes verificadas (2026)")
    canvas.restoreState()


def _page(canvas, doc):
    canvas.saveState(); W, H = A4
    canvas.setFont(F["r"], 8.5); canvas.setFillColor(SUB)
    canvas.drawRightString(W-2.2*cm, H-1.4*cm, "Dossier de Investigación")
    canvas.drawCentredString(W/2, 1.3*cm, str(doc.page))
    canvas.restoreState()


def construir():
    S = estilos()
    out = C.ROOT / "docs" / "Dossier_Investigacion.pdf"
    doc = BaseDocTemplate(str(out), pagesize=A4, leftMargin=2.2*cm, rightMargin=2.2*cm,
                          topMargin=2.2*cm, bottomMargin=2.0*cm, title="Dossier de Investigación")
    fr = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id="m")
    from reportlab.platypus import NextPageTemplate
    doc.addPageTemplates([PageTemplate(id="cover", frames=[fr], onPage=_cover),
                          PageTemplate(id="body", frames=[fr], onPage=_page)])
    flow = [NextPageTemplate("body"), PageBreak()]
    for n, (rel, titulo) in enumerate(DOCS, 1):
        p = C.ROOT / rel
        if not p.exists(): continue
        flow.append(Paragraph(f"{n:02d}", S["eyebrow"]))
        flow.append(Paragraph(titulo, S["h1"]))
        flow += parse_md(p.read_text(encoding="utf-8"), S, skip_h1=True)
        flow.append(PageBreak())
    doc.build(flow)
    print(f"OK {out} ({out.stat().st_size/1024:.0f} KB)")
    _word(S)
    # copiar a web
    import shutil
    for f in ["Dossier_Investigacion.pdf", "Dossier_Investigacion.docx"]:
        s = C.ROOT / "docs" / f; d = C.ROOT / "web" / "assets" / "docs" / f
        if s.exists():
            try: shutil.copy(s, d)
            except Exception as e: print("copia web:", e)
    print("OK copias web")


def _word(S):
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document()
    st = doc.styles["Normal"]; st.font.name = "Segoe UI"; st.font.size = Pt(10.5)
    st.paragraph_format.space_after = Pt(8); st.paragraph_format.line_spacing = 1.25
    NAVYw = RGBColor(0x1B, 0x2A, 0x41); COPPERw = RGBColor(0xC2, 0x70, 0x3D); SUBw = RGBColor(0x6E, 0x6E, 0x73)
    for i, sz, col in [(1, 22, NAVYw), (2, 15, RGBColor(0x1D, 0x1D, 0x1F)), (3, 12, COPPERw)]:
        h = doc.styles[f"Heading {i}"]; h.font.name = "Segoe UI Semibold"; h.font.size = Pt(sz)
        h.font.color.rgb = col; h.font.bold = True
    for s in doc.sections:
        s.top_margin = Cm(2.2); s.bottom_margin = Cm(2.0); s.left_margin = Cm(2.2); s.right_margin = Cm(2.2)
    # portada simple
    t = doc.add_paragraph(); r = t.add_run("Dossier de Investigación"); r.bold = True; r.font.size = Pt(28); r.font.name = "Segoe UI Semibold"; r.font.color.rgb = NAVYw
    s2 = doc.add_paragraph(); r2 = s2.add_run("Estado del arte · Historia del cobre · Bibliografía · Bitácora"); r2.font.size = Pt(12); r2.font.color.rgb = SUBw
    doc.add_paragraph("Material de investigación · Magíster en Data Science · Universidad San Sebastián")
    doc.add_page_break()

    def runs(p, text):
        for part in re.split(r"(\*\*.+?\*\*)", text):
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(re.sub(r"\*\*|`", "", part)); run.bold = True
            else:
                p.add_run(re.sub(r"\*\*|`", "", re.sub(r"\[(.+?)\]\((.+?)\)", r"\1", part)))
    for n, (rel, titulo) in enumerate(DOCS, 1):
        pth = C.ROOT / rel
        if not pth.exists(): continue
        doc.add_heading(titulo, level=1)
        md = pth.read_text(encoding="utf-8"); lineas = md.split("\n"); i = 0; buf = []; skipped = False
        def flush():
            if buf:
                runs(doc.add_paragraph(), " ".join(x.strip() for x in buf)); buf.clear()
        while i < len(lineas):
            s = lineas[i].strip()
            if not s: flush(); i += 1; continue
            if s.startswith("|"):
                flush()
                while i < len(lineas) and lineas[i].strip().startswith("|"): i += 1
                continue
            if s.startswith("# "):
                flush()
                if not skipped: skipped = True; i += 1; continue
                doc.add_heading(re.sub(r"`", "", s[2:]), level=2)
            elif s.startswith("### "): flush(); doc.add_heading(s[4:], level=3)
            elif s.startswith("## "): flush(); doc.add_heading(s[3:], level=2)
            elif re.match(r"^[-*] ", s): flush(); runs(doc.add_paragraph(style="List Bullet"), s[2:])
            elif s.startswith("> "): flush(); runs(doc.add_paragraph(style="Intense Quote"), s[2:])
            elif set(s) <= set("-") and len(s) >= 3: flush()
            else: buf.append(s)
            i += 1
        flush(); doc.add_page_break()
    outw = C.ROOT / "docs" / "Dossier_Investigacion.docx"
    try:
        doc.save(outw); print(f"OK {outw}")
    except PermissionError:
        print("AVISO Dossier docx bloqueado")


if __name__ == "__main__":
    construir()
